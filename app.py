"""
Meeting Summarizer & Action Item Extractor
Author: Rishi Raj
Stack: Python + Streamlit + Whisper + Claude API
UI: Cinematic Classic Light — Cormorant serif, cream palette, 3D depth
"""

import os
import json
import tempfile
import traceback
from pathlib import Path

import streamlit as st
from dotenv import load_dotenv

load_dotenv()

# ─────────────────────────── Page config ───────────────────────────
st.set_page_config(
    page_title="Meeting Summarizer",
    page_icon="🎬",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─────────────────────────── CINEMATIC CSS ──────────────────────────
st.markdown("""
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Cormorant+Garamond:ital,wght@0,300;0,400;0,500;0,600;1,300;1,400&family=DM+Sans:wght@300;400;500&family=Playfair+Display:wght@400;700&display=swap" rel="stylesheet">

<style>
/* ── Root variables ── */
:root {
  --cream:       #F7F4EE;
  --cream-dark:  #EDE9DF;
  --paper:       #FDFCFA;
  --ink:         #1C1917;
  --ink-muted:   #6B6560;
  --gold:        #B8964E;
  --gold-light:  #D4AF70;
  --gold-pale:   #F0E6CC;
  --white:       #FFFFFF;
  --shadow-sm:   0 2px 8px rgba(28,25,23,0.06), 0 1px 3px rgba(28,25,23,0.04);
  --shadow-md:   0 8px 24px rgba(28,25,23,0.10), 0 3px 8px rgba(28,25,23,0.06), 0 1px 2px rgba(28,25,23,0.04);
  --shadow-lg:   0 20px 60px rgba(28,25,23,0.14), 0 8px 20px rgba(28,25,23,0.08), 0 2px 6px rgba(28,25,23,0.05);
  --shadow-3d:   0 1px 0 rgba(255,255,255,0.9) inset,
                 0 24px 48px rgba(28,25,23,0.12),
                 0 8px 16px rgba(28,25,23,0.08),
                 0 2px 4px rgba(28,25,23,0.06);
  --border:      rgba(28,25,23,0.08);
  --radius:      14px;
  --radius-lg:   20px;
}

/* ── Global reset ── */
html, body, [class*="css"] {
  font-family: 'DM Sans', sans-serif !important;
  background-color: var(--cream) !important;
  color: var(--ink) !important;
}

/* ── Grain texture overlay ── */
body::before {
  content: '';
  position: fixed;
  inset: 0;
  background-image: url("data:image/svg+xml,%3Csvg viewBox='0 0 256 256' xmlns='http://www.w3.org/2000/svg'%3E%3Cfilter id='noise'%3E%3CfeTurbulence type='fractalNoise' baseFrequency='0.9' numOctaves='4' stitchTiles='stitch'/%3E%3C/filter%3E%3Crect width='100%25' height='100%25' filter='url(%23noise)' opacity='0.03'/%3E%3C/svg%3E");
  pointer-events: none;
  z-index: 9999;
  opacity: 0.4;
}

/* ── Streamlit chrome cleanup ── */
#MainMenu, footer, header { visibility: hidden !important; }
.block-container { 
  padding: 2rem 3rem 4rem !important; 
  max-width: 1200px !important;
}
.stApp { background-color: var(--cream) !important; }
section[data-testid="stSidebar"] > div {
  background: var(--paper) !important;
  border-right: 1px solid var(--border) !important;
  box-shadow: 4px 0 24px rgba(28,25,23,0.05) !important;
}

/* ── Sidebar content ── */
section[data-testid="stSidebar"] .stSelectbox label,
section[data-testid="stSidebar"] .stMarkdown p,
section[data-testid="stSidebar"] h1,
section[data-testid="stSidebar"] h2,
section[data-testid="stSidebar"] h3 {
  font-family: 'DM Sans', sans-serif !important;
  color: var(--ink) !important;
}

/* ── Selectbox ── */
.stSelectbox > div > div {
  background: var(--white) !important;
  border: 1px solid var(--border) !important;
  border-radius: 8px !important;
  box-shadow: var(--shadow-sm) !important;
  font-family: 'DM Sans', sans-serif !important;
  font-size: 0.88rem !important;
  color: var(--ink) !important;
  transition: box-shadow 0.2s ease, border-color 0.2s ease !important;
}
.stSelectbox > div > div:hover {
  border-color: var(--gold) !important;
  box-shadow: 0 0 0 3px var(--gold-pale) !important;
}

/* ── File uploader ── */
[data-testid="stFileUploader"] {
  background: var(--white) !important;
  border: 2px dashed var(--cream-dark) !important;
  border-radius: var(--radius-lg) !important;
  padding: 2.5rem !important;
  transition: all 0.3s ease !important;
  box-shadow: var(--shadow-3d) !important;
  position: relative !important;
}
[data-testid="stFileUploader"]:hover {
  border-color: var(--gold-light) !important;
  box-shadow: 0 1px 0 rgba(255,255,255,0.9) inset,
              0 28px 56px rgba(28,25,23,0.14),
              0 10px 20px rgba(28,25,23,0.08),
              0 0 0 4px var(--gold-pale) !important;
  transform: translateY(-2px) !important;
}
[data-testid="stFileUploader"] label {
  font-family: 'DM Sans', sans-serif !important;
  font-size: 0.9rem !important;
  color: var(--ink-muted) !important;
}
[data-testid="stFileUploaderDropzone"] {
  background: transparent !important;
  border: none !important;
}

/* ── Primary button ── */
.stButton > button[kind="primary"] {
  background: linear-gradient(160deg, #2A2520 0%, #1C1917 100%) !important;
  color: #F0E6CC !important;
  font-family: 'DM Sans', sans-serif !important;
  font-weight: 500 !important;
  letter-spacing: 0.08em !important;
  font-size: 0.82rem !important;
  text-transform: uppercase !important;
  border: none !important;
  border-radius: 10px !important;
  padding: 0.9rem 2rem !important;
  box-shadow: 0 1px 0 rgba(255,255,255,0.1) inset,
              0 12px 32px rgba(28,25,23,0.3),
              0 4px 10px rgba(28,25,23,0.15),
              0 1px 3px rgba(28,25,23,0.1) !important;
  transition: all 0.25s cubic-bezier(0.34, 1.56, 0.64, 1) !important;
  cursor: pointer !important;
}
.stButton > button[kind="primary"]:hover {
  transform: translateY(-3px) scale(1.01) !important;
  box-shadow: 0 1px 0 rgba(255,255,255,0.12) inset,
              0 20px 48px rgba(28,25,23,0.35),
              0 8px 16px rgba(28,25,23,0.18),
              0 2px 4px rgba(28,25,23,0.1) !important;
  background: linear-gradient(160deg, #3A3028 0%, #2A2218 100%) !important;
}
.stButton > button[kind="primary"]:active {
  transform: translateY(0px) scale(0.99) !important;
}

/* ── Secondary / download button ── */
.stDownloadButton > button {
  background: var(--white) !important;
  color: var(--ink) !important;
  font-family: 'DM Sans', sans-serif !important;
  font-weight: 500 !important;
  letter-spacing: 0.05em !important;
  font-size: 0.82rem !important;
  text-transform: uppercase !important;
  border: 1px solid var(--border) !important;
  border-radius: 10px !important;
  padding: 0.8rem 2rem !important;
  box-shadow: var(--shadow-3d) !important;
  transition: all 0.25s cubic-bezier(0.34, 1.56, 0.64, 1) !important;
}
.stDownloadButton > button:hover {
  border-color: var(--gold) !important;
  color: var(--gold) !important;
  transform: translateY(-2px) !important;
  box-shadow: 0 1px 0 rgba(255,255,255,0.9) inset,
              0 16px 36px rgba(28,25,23,0.12),
              0 0 0 3px var(--gold-pale) !important;
}

/* ── Expander ── */
.streamlit-expanderHeader {
  background: var(--white) !important;
  border: 1px solid var(--border) !important;
  border-radius: var(--radius) !important;
  font-family: 'DM Sans', sans-serif !important;
  font-size: 0.88rem !important;
  font-weight: 500 !important;
  color: var(--ink) !important;
  padding: 0.85rem 1.2rem !important;
  box-shadow: var(--shadow-sm) !important;
  transition: all 0.2s ease !important;
}
.streamlit-expanderHeader:hover {
  border-color: var(--gold-light) !important;
  box-shadow: var(--shadow-md) !important;
}
.streamlit-expanderContent {
  background: var(--paper) !important;
  border: 1px solid var(--border) !important;
  border-top: none !important;
  border-radius: 0 0 var(--radius) var(--radius) !important;
  padding: 1rem !important;
}

/* ── Text areas ── */
.stTextArea textarea {
  background: var(--paper) !important;
  border: 1px solid var(--border) !important;
  border-radius: 8px !important;
  font-family: 'DM Sans', sans-serif !important;
  font-size: 0.84rem !important;
  color: var(--ink) !important;
  line-height: 1.7 !important;
}

/* ── Metrics ── */
[data-testid="metric-container"] {
  background: var(--white) !important;
  border: 1px solid var(--border) !important;
  border-radius: var(--radius) !important;
  padding: 1.4rem 1.6rem !important;
  box-shadow: var(--shadow-3d) !important;
  transition: transform 0.2s ease, box-shadow 0.2s ease !important;
}
[data-testid="metric-container"]:hover {
  transform: translateY(-3px) !important;
  box-shadow: 0 1px 0 rgba(255,255,255,0.9) inset,
              0 24px 48px rgba(28,25,23,0.14),
              0 8px 16px rgba(28,25,23,0.08) !important;
}
[data-testid="metric-container"] label {
  font-family: 'DM Sans', sans-serif !important;
  font-size: 0.75rem !important;
  font-weight: 500 !important;
  letter-spacing: 0.1em !important;
  text-transform: uppercase !important;
  color: var(--ink-muted) !important;
}
[data-testid="metric-container"] [data-testid="stMetricValue"] {
  font-family: 'Cormorant Garamond', serif !important;
  font-size: 2.4rem !important;
  font-weight: 600 !important;
  color: var(--ink) !important;
  line-height: 1.1 !important;
}

/* ── Status / spinner ── */
[data-testid="stStatusWidget"] {
  background: var(--white) !important;
  border: 1px solid var(--border) !important;
  border-radius: var(--radius) !important;
  box-shadow: var(--shadow-md) !important;
}

/* ── Divider ── */
hr {
  border: none !important;
  border-top: 1px solid var(--border) !important;
  margin: 2rem 0 !important;
}

/* ── Caption / small text ── */
.stCaption, small {
  font-family: 'DM Sans', sans-serif !important;
  font-size: 0.78rem !important;
  color: var(--ink-muted) !important;
  letter-spacing: 0.02em !important;
}

/* ── Alert/Info boxes ── */
[data-testid="stAlert"] {
  border-radius: var(--radius) !important;
  border: 1px solid var(--border) !important;
  font-family: 'DM Sans', sans-serif !important;
  font-size: 0.88rem !important;
  box-shadow: var(--shadow-sm) !important;
}

/* ── Scrollbar ── */
::-webkit-scrollbar { width: 6px; height: 6px; }
::-webkit-scrollbar-track { background: var(--cream); }
::-webkit-scrollbar-thumb { background: var(--cream-dark); border-radius: 3px; }
::-webkit-scrollbar-thumb:hover { background: var(--gold-light); }

/* ── Animations ── */
@keyframes fadeUp {
  from { opacity: 0; transform: translateY(20px); }
  to   { opacity: 1; transform: translateY(0); }
}
@keyframes shimmer {
  0%   { background-position: -200% center; }
  100% { background-position: 200% center; }
}
.fade-up { animation: fadeUp 0.6s cubic-bezier(0.22, 1, 0.36, 1) both; }
.fade-up-2 { animation: fadeUp 0.6s 0.1s cubic-bezier(0.22, 1, 0.36, 1) both; }
.fade-up-3 { animation: fadeUp 0.6s 0.2s cubic-bezier(0.22, 1, 0.36, 1) both; }

</style>
""", unsafe_allow_html=True)


# ─────────────────────────── Lazy imports ───────────────────────────
@st.cache_resource
def load_whisper_model(model_size="base"):
    import whisper
    return whisper.load_model(model_size)

def import_pdfplumber():
    import pdfplumber
    return pdfplumber

def import_docx():
    from docx import Document
    return Document

def import_pillow_tesseract():
    from PIL import Image
    import pytesseract
    return Image, pytesseract

def get_anthropic_client():
    import anthropic
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        st.error("❌ ANTHROPIC_API_KEY not set. Add it to your .env file or environment.")
        st.stop()
    return anthropic.Anthropic(api_key=api_key)


# ─────────────────────────── Extraction helpers ─────────────────────
def extract_audio(file_bytes, filename):
    suffix = Path(filename).suffix.lower()
    try:
        model = load_whisper_model(st.session_state.get("whisper_model", "base"))
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        result = model.transcribe(tmp_path)
        os.unlink(tmp_path)
        return result["text"].strip(), None
    except Exception as e:
        return "", f"Whisper error: {e}\n\nTip: Check that ffmpeg is installed and the audio format is supported."

def extract_pdf(file_bytes, filename):
    try:
        pdfplumber = import_pdfplumber()
        import io
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"[Page {i}]\n{page_text}")
        if not text_parts:
            return "", "PDF appears to have no extractable text (possibly scanned)."
        return "\n\n".join(text_parts), None
    except Exception as e:
        return "", f"PDF extraction error: {e}"

def extract_docx(file_bytes, filename):
    try:
        import io
        Document = import_docx()
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        if not paragraphs:
            return "", "DOCX has no readable text."
        return "\n".join(paragraphs), None
    except Exception as e:
        return "", f"DOCX extraction error: {e}"

def extract_txt(file_bytes, filename):
    try:
        return file_bytes.decode("utf-8", errors="replace").strip(), None
    except Exception as e:
        return "", f"Text file read error: {e}"

def extract_image(file_bytes, filename):
    try:
        import io
        Image, pytesseract = import_pillow_tesseract()
        img = Image.open(io.BytesIO(file_bytes)).convert("L")
        text = pytesseract.image_to_string(img)
        if not text.strip():
            return "", "No text detected in image."
        return text.strip(), None
    except ImportError:
        return "", (
            "OCR requires **Tesseract** to be installed.\n\n"
            "• macOS: `brew install tesseract`\n"
            "• Ubuntu: `sudo apt-get install tesseract-ocr`\n"
            "• Windows: https://github.com/UB-Mannheim/tesseract/wiki"
        )
    except Exception as e:
        return "", f"Image OCR error: {e}"

EXTRACTOR_MAP = {
    ".mp3": extract_audio, ".wav": extract_audio, ".m4a": extract_audio,
    ".flac": extract_audio, ".ogg": extract_audio,
    ".pdf":  extract_pdf,
    ".docx": extract_docx,
    ".txt":  extract_txt,
    ".png":  extract_image, ".jpg": extract_image, ".jpeg": extract_image,
}
AUDIO_EXTS = {".mp3", ".wav", ".m4a", ".flac", ".ogg"}

def process_file(uploaded_file):
    name = uploaded_file.name
    ext  = Path(name).suffix.lower()
    file_bytes = uploaded_file.read()
    extractor = EXTRACTOR_MAP.get(ext)
    if extractor is None:
        return {"name": name, "ext": ext, "text": "", "error": f"Unsupported file type: {ext}", "ok": False}
    text, error = extractor(file_bytes, name)
    return {"name": name, "ext": ext, "text": text, "error": error, "ok": bool(text and not error)}


# ─────────────────────────── Claude integration ─────────────────────
SYSTEM_PROMPT = """You are an expert meeting analyst. You will receive combined text from multiple meeting artifacts.

Return ONLY a JSON object — no markdown, no preamble:

{
  "summary": "<concise synthesis, max 200 words, highlight key decisions>",
  "action_items": [
    {"task": "<clear action>", "responsible": "<person or TBD>", "due": "<deadline or TBD>"}
  ]
}

Rules: summary ≤200 words. Extract EVERY action item. Empty list if none."""

def call_claude(combined_text, model):
    client = get_anthropic_client()
    try:
        message = client.messages.create(
            model=model, max_tokens=1500,
            system=SYSTEM_PROMPT,
            messages=[{"role": "user", "content":
                "Here are the meeting artifacts:\n\n" + combined_text + "\n\nGenerate the JSON output now."}],
        )
        raw = message.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        return json.loads(raw), None
    except json.JSONDecodeError:
        return None, f"Claude returned non-JSON:\n\n```\n{raw}\n```"
    except Exception as e:
        return None, f"Claude API error: {e}\n\n{traceback.format_exc()}"

def build_markdown(results, claude_output):
    lines = ["# Meeting Summary & Action Items\n",
             "## Summary\n", claude_output.get("summary", ""),
             "\n\n## Action Items\n"]
    items = claude_output.get("action_items", [])
    if items:
        lines += ["| # | Task | Responsible | Due |", "|---|------|-------------|-----|"]
        for i, item in enumerate(items, 1):
            lines.append(f"| {i} | {item.get('task','—')} | {item.get('responsible','TBD')} | {item.get('due','TBD')} |")
    else:
        lines.append("No action items identified.")
    lines.append("\n\n## Source Files\n")
    for r in results:
        lines.append(f"- {'✅' if r['ok'] else '❌'} `{r['name']}`")
    return "\n".join(lines)


# ─────────────────────────── Session state ──────────────────────────
def init_state():
    for k, v in [("results", []), ("claude_output", None), ("whisper_model", "base")]:
        if k not in st.session_state:
            st.session_state[k] = v


# ─────────────────────────── UI Components ──────────────────────────

def render_header():
    st.markdown("""
    <div class="fade-up" style="
        text-align: center;
        padding: 4rem 2rem 3rem;
        position: relative;
    ">
        <!-- Decorative line top -->
        <div style="
            display: flex; align-items: center; justify-content: center;
            gap: 1rem; margin-bottom: 1.8rem;
        ">
            <div style="height:1px; width:60px; background: linear-gradient(90deg, transparent, var(--gold));"></div>
            <span style="font-family:'DM Sans',sans-serif; font-size:0.7rem; letter-spacing:0.25em;
                         text-transform:uppercase; color:var(--gold); font-weight:500;">
                Est. 2025
            </span>
            <div style="height:1px; width:60px; background: linear-gradient(90deg, var(--gold), transparent);"></div>
        </div>

        <!-- Main title -->
        <h1 style="
            font-family: 'Cormorant Garamond', serif;
            font-size: clamp(2.8rem, 6vw, 5rem);
            font-weight: 300;
            letter-spacing: -0.01em;
            line-height: 1.05;
            color: var(--ink);
            margin: 0 0 0.5rem;
        ">
            Meeting<br>
            <em style="font-style:italic; font-weight:400; color:#4A3728;">Summarizer</em>
        </h1>

        <!-- Subtitle -->
        <p style="
            font-family: 'DM Sans', sans-serif;
            font-size: 0.85rem;
            letter-spacing: 0.18em;
            text-transform: uppercase;
            color: var(--ink-muted);
            font-weight: 400;
            margin: 1.2rem 0 0;
        ">
            Action Items &nbsp;·&nbsp; Unified Summary &nbsp;·&nbsp; All File Types
        </p>

        <!-- Decorative bottom ornament -->
        <div style="margin-top: 2rem; color: var(--gold); font-size: 1.2rem; letter-spacing: 0.5em;">
            ✦ ✦ ✦
        </div>
    </div>
    """, unsafe_allow_html=True)


def render_sidebar():
    with st.sidebar:
        st.markdown("""
        <div style="padding: 1.5rem 0 1rem;">
            <p style="font-family:'Cormorant Garamond',serif; font-size:1.4rem;
                      font-weight:600; color:var(--ink); margin:0 0 0.2rem;">
                Settings
            </p>
            <div style="height:2px; width:32px;
                        background:linear-gradient(90deg,var(--gold),transparent);
                        margin-bottom:1.5rem;"></div>
        </div>
        """, unsafe_allow_html=True)

        st.markdown('<p style="font-size:0.72rem;letter-spacing:0.12em;text-transform:uppercase;color:var(--ink-muted);font-weight:500;margin-bottom:0.3rem;">Whisper Model</p>', unsafe_allow_html=True)
        whisper_model = st.selectbox("Whisper model", ["tiny","base","small","medium","large"],
                                      index=1, label_visibility="collapsed")
        st.session_state.whisper_model = whisper_model

        st.markdown('<p style="font-size:0.72rem;letter-spacing:0.12em;text-transform:uppercase;color:var(--ink-muted);font-weight:500;margin-top:1rem;margin-bottom:0.3rem;">Claude Model</p>', unsafe_allow_html=True)
        claude_model = st.selectbox("Claude model",
            ["claude-sonnet-4-20250514", "claude-haiku-4-5-20251001"],
            index=0, label_visibility="collapsed")

        st.markdown("<hr>", unsafe_allow_html=True)

        st.markdown("""
        <div style="padding: 0.5rem 0;">
            <p style="font-size:0.72rem;letter-spacing:0.12em;text-transform:uppercase;
                      color:var(--ink-muted);font-weight:500;margin-bottom:0.8rem;">
                Accepted Formats
            </p>
            <div style="display:flex;flex-direction:column;gap:0.5rem;">
        """ + "".join([
            f'<div style="display:flex;align-items:center;gap:0.6rem;">'
            f'<span style="font-size:0.95rem;">{icon}</span>'
            f'<span style="font-family:DM Sans,sans-serif;font-size:0.82rem;color:var(--ink-muted);">{label}</span>'
            f'</div>'
            for icon, label in [
                ("🎵", "Audio — mp3, wav, m4a, flac"),
                ("📄", "PDF — pdf"),
                ("📝", "Word — docx"),
                ("📃", "Text — txt"),
                ("🖼️", "Image — png, jpg, jpeg"),
            ]
        ]) + """
            </div>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("<hr>", unsafe_allow_html=True)
        st.markdown("""
        <p style="font-family:'Cormorant Garamond',serif;font-style:italic;
                  font-size:0.95rem;color:var(--ink-muted);text-align:center;margin:0;">
            Built by Rishi Raj
        </p>
        """, unsafe_allow_html=True)

    return claude_model


def render_uploader():
    st.markdown("""
    <div class="fade-up-2" style="margin-bottom:0.5rem;">
        <p style="font-family:'DM Sans',sans-serif;font-size:0.72rem;letter-spacing:0.15em;
                  text-transform:uppercase;color:var(--ink-muted);font-weight:500;margin-bottom:0.6rem;">
            Upload Meeting Files
        </p>
    </div>
    """, unsafe_allow_html=True)

    uploaded_files = st.file_uploader(
        label="Drop files here",
        accept_multiple_files=True,
        type=["mp3","wav","m4a","flac","ogg","pdf","docx","txt","png","jpg","jpeg"],
        label_visibility="collapsed",
    )
    return uploaded_files


def render_file_card(r):
    ok = r["ok"]
    ext = r["ext"].upper().lstrip(".")
    icon_map = {
        ".mp3":"🎵",".wav":"🎵",".m4a":"🎵",".flac":"🎵",".ogg":"🎵",
        ".pdf":"📄",".docx":"📝",".txt":"📃",".png":"🖼️",".jpg":"🖼️",".jpeg":"🖼️",
    }
    file_icon = icon_map.get(r["ext"], "📎")
    status_color = "#27AE60" if ok else "#C0392B"
    status_dot   = "●" if ok else "●"

    with st.expander(f"{file_icon}  {r['name']}  ·  {ext}", expanded=False):
        if r["error"]:
            st.error(r["error"])
        if r["text"]:
            word_count = len(r["text"].split())
            st.markdown(f'<p style="font-size:0.75rem;letter-spacing:0.08em;text-transform:uppercase;color:var(--ink-muted);margin-bottom:0.5rem;">~{word_count:,} words extracted</p>', unsafe_allow_html=True)
            st.text_area("text", value=r["text"], height=180, disabled=True,
                         key=f"ta_{r['name']}", label_visibility="collapsed")


def render_results(results, claude_output):
    # ── Metrics row ──
    n_sources = sum(1 for r in results if r["ok"])
    n_words   = len(claude_output.get("summary","").split())
    n_actions = len(claude_output.get("action_items",[]))

    st.markdown('<div class="fade-up" style="margin:2rem 0 1rem;">', unsafe_allow_html=True)
    c1, c2, c3 = st.columns(3)
    with c1: st.metric("Sources Processed", n_sources)
    with c2: st.metric("Summary Words", n_words)
    with c3: st.metric("Action Items", n_actions)
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown("<hr>", unsafe_allow_html=True)

    # ── Summary ──
    st.markdown("""
    <div class="fade-up" style="margin-bottom:0.6rem;">
        <p style="font-family:'DM Sans',sans-serif;font-size:0.7rem;letter-spacing:0.2em;
                  text-transform:uppercase;color:var(--gold);font-weight:500;">
            Meeting Summary
        </p>
        <p style="font-family:'Cormorant Garamond',serif;font-size:1.7rem;font-weight:400;
                  color:var(--ink);margin:0 0 1rem;line-height:1.2;">
            What was discussed
        </p>
    </div>
    """, unsafe_allow_html=True)

    summary_text = claude_output.get("summary", "No summary generated.")
    st.markdown(f"""
    <div class="fade-up-2" style="
        background: var(--white);
        border: 1px solid var(--border);
        border-radius: var(--radius-lg);
        padding: 2.2rem 2.5rem;
        box-shadow: var(--shadow-3d);
        position: relative;
        overflow: hidden;
        margin-bottom: 2rem;
    ">
        <!-- Gold accent bar -->
        <div style="position:absolute;top:0;left:0;bottom:0;width:4px;
                    background:linear-gradient(180deg,var(--gold-light),var(--gold));
                    border-radius:4px 0 0 4px;"></div>
        <!-- Quote mark -->
        <div style="font-family:'Cormorant Garamond',serif;font-size:5rem;
                    color:var(--gold-pale);line-height:1;position:absolute;
                    top:0.5rem;right:1.5rem;font-weight:700;user-select:none;">"</div>
        <p style="font-family:'Cormorant Garamond',serif;font-size:1.18rem;
                  line-height:1.85;color:var(--ink);margin:0;font-weight:400;
                  font-style:italic;">
            {summary_text}
        </p>
    </div>
    """, unsafe_allow_html=True)

    # ── Action Items ──
    st.markdown("""
    <div class="fade-up" style="margin-bottom:1rem;">
        <p style="font-family:'DM Sans',sans-serif;font-size:0.7rem;letter-spacing:0.2em;
                  text-transform:uppercase;color:var(--gold);font-weight:500;">
            Action Items
        </p>
        <p style="font-family:'Cormorant Garamond',serif;font-size:1.7rem;font-weight:400;
                  color:var(--ink);margin:0 0 1rem;line-height:1.2;">
            Next steps &amp; owners
        </p>
    </div>
    """, unsafe_allow_html=True)

    items = claude_output.get("action_items", [])
    if items:
        for i, item in enumerate(items, 1):
            task        = item.get("task", "—")
            responsible = item.get("responsible", "TBD")
            due         = item.get("due", "TBD")
            st.markdown(f"""
            <div class="fade-up" style="
                background: var(--white);
                border: 1px solid var(--border);
                border-radius: var(--radius);
                padding: 1.4rem 1.8rem;
                margin-bottom: 0.75rem;
                box-shadow: var(--shadow-3d);
                display: flex;
                align-items: flex-start;
                gap: 1.4rem;
                transition: transform 0.2s ease, box-shadow 0.2s ease;
            " onmouseover="this.style.transform='translateY(-3px)';this.style.boxShadow='0 1px 0 rgba(255,255,255,0.9) inset, 0 24px 48px rgba(28,25,23,0.14), 0 8px 16px rgba(28,25,23,0.08)'"
               onmouseout="this.style.transform='';this.style.boxShadow='var(--shadow-3d)'">
                <!-- Number badge -->
                <div style="
                    min-width: 2.4rem; height: 2.4rem;
                    background: linear-gradient(135deg, #2A2520, #1C1917);
                    border-radius: 50%;
                    display: flex; align-items: center; justify-content: center;
                    box-shadow: 0 4px 12px rgba(28,25,23,0.25);
                    flex-shrink: 0; margin-top: 0.1rem;
                ">
                    <span style="font-family:'Cormorant Garamond',serif;font-size:1rem;
                                 font-weight:600;color:#F0E6CC;">{i}</span>
                </div>
                <!-- Content -->
                <div style="flex:1;">
                    <p style="font-family:'Cormorant Garamond',serif;font-size:1.12rem;
                               font-weight:500;color:var(--ink);margin:0 0 0.6rem;
                               line-height:1.4;">{task}</p>
                    <div style="display:flex;gap:0.75rem;flex-wrap:wrap;">
                        <span style="
                            font-family:'DM Sans',sans-serif;font-size:0.75rem;
                            letter-spacing:0.05em;font-weight:500;
                            background:var(--gold-pale);color:#7A5C20;
                            padding:0.25rem 0.75rem;border-radius:20px;
                            border:1px solid rgba(184,150,78,0.3);
                        ">👤 {responsible}</span>
                        <span style="
                            font-family:'DM Sans',sans-serif;font-size:0.75rem;
                            letter-spacing:0.05em;font-weight:500;
                            background:#EEF0FF;color:#4040A0;
                            padding:0.25rem 0.75rem;border-radius:20px;
                            border:1px solid rgba(64,64,160,0.2);
                        ">📅 {due}</span>
                    </div>
                </div>
            </div>
            """, unsafe_allow_html=True)
    else:
        st.info("No action items were identified in the uploaded content.")

    st.markdown("<hr>", unsafe_allow_html=True)

    # ── Source files ──
    st.markdown("""
    <p style="font-family:'DM Sans',sans-serif;font-size:0.7rem;letter-spacing:0.2em;
              text-transform:uppercase;color:var(--ink-muted);font-weight:500;margin-bottom:1rem;">
        Source Files
    </p>
    """, unsafe_allow_html=True)
    render_file_results(results)

    st.markdown("<div style='height:1.5rem;'></div>", unsafe_allow_html=True)

    # ── Download ──
    md = build_markdown(results, claude_output)
    st.download_button(
        label="⬇   Download as Markdown",
        data=md,
        file_name="meeting_summary.md",
        mime="text/markdown",
        use_container_width=True,
    )


def render_file_results(results):
    for r in results:
        render_file_card(r)


def render_empty_state():
    st.markdown("""
    <div class="fade-up" style="
        text-align: center;
        padding: 4rem 2rem;
        background: var(--white);
        border: 1px solid var(--border);
        border-radius: var(--radius-lg);
        box-shadow: var(--shadow-3d);
        margin-top: 1rem;
    ">
        <div style="font-size:3.5rem;margin-bottom:1.5rem;opacity:0.6;">🎬</div>
        <p style="font-family:'Cormorant Garamond',serif;font-size:1.6rem;
                  font-weight:400;color:var(--ink);margin:0 0 0.75rem;font-style:italic;">
            Ready for your meeting files
        </p>
        <p style="font-family:'DM Sans',sans-serif;font-size:0.84rem;
                  color:var(--ink-muted);max-width:380px;margin:0 auto;line-height:1.7;">
            Upload any combination of audio, PDF, Word, text, or image files above.<br>
            Claude will synthesize everything into a clean summary.
        </p>
        <div style="margin-top:2rem;display:flex;justify-content:center;gap:1rem;flex-wrap:wrap;">
    """ + "".join([
        f'<span style="font-family:DM Sans,sans-serif;font-size:0.75rem;letter-spacing:0.08em;'
        f'text-transform:uppercase;color:var(--ink-muted);background:var(--cream);'
        f'padding:0.4rem 1rem;border-radius:20px;border:1px solid var(--border);">{t}</span>'
        for t in ["Audio", "PDF", "Word", "Images", "Text"]
    ]) + """
        </div>
    </div>
    """, unsafe_allow_html=True)


# ─────────────────────────── Main ───────────────────────────────────
def main():
    init_state()
    render_header()
    claude_model = render_sidebar()

    # ── Upload zone ──
    col_upload, col_gap = st.columns([3, 1])
    with col_upload:
        uploaded_files = render_uploader()

    if uploaded_files:
        st.markdown(f"""
        <p style="font-family:'DM Sans',sans-serif;font-size:0.8rem;
                  color:var(--ink-muted);margin:0.5rem 0 1rem;letter-spacing:0.04em;">
            {len(uploaded_files)} file{'s' if len(uploaded_files)>1 else ''} selected
        </p>
        """, unsafe_allow_html=True)

        if st.button("  Process & Summarize  ", type="primary", use_container_width=False):
            st.session_state.results      = []
            st.session_state.claude_output = None

            has_audio = any(Path(f.name).suffix.lower() in AUDIO_EXTS for f in uploaded_files)
            label = "Transcribing audio & extracting text…" if has_audio else "Extracting text from files…"

            with st.status(label, expanded=True) as status:
                results = []
                for uf in uploaded_files:
                    st.write(f"Processing `{uf.name}`…")
                    result = process_file(uf)
                    results.append(result)
                    if result["ok"]:
                        st.write(f"  ✅  ~{len(result['text'].split()):,} words extracted")
                    else:
                        st.write(f"  ⚠️  {result['error']}")

                st.session_state.results = results
                good = [r for r in results if r["ok"]]

                if not good:
                    status.update(label="❌ No text extracted", state="error")
                    st.error("Could not extract text from any file. Check the errors above.")
                    return

                combined = "\n\n" + "="*60 + "\n\n".join(
                    f"[Source: {r['name']}]\n{r['text']}" for r in good
                )
                total_words = len(combined.split())
                if total_words > 6000:
                    st.warning(f"⚠️ Combined text is ~{total_words:,} words. Consider fewer files for best results.")

                st.write("🤖  Sending to Claude for analysis…")
                out, err = call_claude(combined, claude_model)
                if err:
                    status.update(label="⚠️ Claude error", state="error")
                    st.error(err)
                    return

                st.session_state.claude_output = out
                status.update(label="✅  Analysis complete", state="complete", expanded=False)

    # ── Render persisted results ──
    if st.session_state.claude_output and st.session_state.results:
        render_results(st.session_state.results, st.session_state.claude_output)
    elif not uploaded_files and not st.session_state.results:
        render_empty_state()


if __name__ == "__main__":
    main()
